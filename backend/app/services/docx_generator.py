from docx import Document
from docx.shared import Pt
import io

def create_docx(resume_content: str) -> bytes:
    doc = Document()
    
    # Very basic styling for the document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    
    for line in resume_content.split('\n'):
        if line.strip() == "":
            doc.add_paragraph("")
            continue
            
        p = doc.add_paragraph()
        
        # Simple markdown handling for bold **text**
        parts = line.split("**")
        for i, part in enumerate(parts):
            if i % 2 == 1:
                # Inside **
                r = p.add_run(part)
                r.bold = True
            else:
                # Outside **
                # handle single * for bullets
                clean_part = part.replace("* ", "• ")
                p.add_run(clean_part)
                
    file_stream = io.BytesIO()
    doc.save(file_stream)
    return file_stream.getvalue()
